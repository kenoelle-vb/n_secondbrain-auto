# streamlit run "C:/Users/keno/OneDrive/Documents/Projects/n_secondbrain applier/Z_TEST2.py"

import streamlit as st
import requests
from bs4 import BeautifulSoup, ParserRejectedMarkup
from newspaper import Article
import google.generativeai as genai
import pandas as pd
import datetime
import time
import threading
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import matplotlib.pyplot as plt
from io import BytesIO
import zipfile
from docx import Document
from PIL import Image
import base64
import random
from urllib.parse import quote_plus
import nltk
from streamlit.components.v1 import html
nltk.download('punkt_tab')

# ============================================== HTML AND CSS =========================================================================================

logo = Image.open("logo.png")

st.set_page_config(
    page_title="n_secondbrain",
    page_icon=logo,
    layout="wide",  # or "wide" if you prefer
    initial_sidebar_state="expanded",
)

st.set_option('client.showErrorDetails', False)

st.markdown(
    """
    <style>
    section[data-testid="stMain"] > div[data-testid="stMainBlockContainer"] {
         padding-top: 0px;  # Remove padding completely
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Hide Streamlit style elements
hide_st_style = """
                <style>
                #MainMenu {visibility: hidden;}
                footer {visibility: hidden;}
                header {visibility: hidden;}
                </style>
                """
st.markdown(hide_st_style, unsafe_allow_html=True)

st.markdown("""
    <style>
    [data-testid="stTextArea"] {
        color: #FFFFFF;
    }
    </style>
    """, unsafe_allow_html=True)

# Set Montserrat font
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;700&display=swap');

    html, body, [class*="css"] {
        font-family: 'Montserrat', sans-serif;
    }
    </style>
    """, unsafe_allow_html=True)

# Change color of specific Streamlit elements
st.markdown("""
    <style>
    .st-emotion-cache-1o6s5t7 {
        color: #ababab !important;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stExpander {
        background-color: #FFFFFF;
        border-radius: 10px;
    }
    
    .stExpander > details {
        background-color: #FFFFFF;
        border-radius: 10px;
    }
    
    .stExpander > details > summary {
        background-color: #FFFFFF;
        border-radius: 10px 10px 0 0;
        padding: 10px;
    }
    
    .stExpander > details > div {
        background-color: #FFFFFF;
        border-radius: 0 0 10px 10px;
        padding: 10px;
    }
    
    .stCheckbox {
        background-color: #FFFFFF;
        border-radius: 5px;
        padding: 5px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown("""
    <style>
    .stButton > button {
        color: #FFFFFF;
        background-color: #424040;
        border-radius: 10px;
    }
    </style>
    """, unsafe_allow_html=True)

st.markdown(
    """
    <style>
    .streamlit-expanderHeader {
        font-size: 20px;
    }
    .streamlit-expanderContent {
        max-height: 400px;
        overflow-y: auto;
    }
    </style>
    """,
    unsafe_allow_html=True
)

def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

def set_png_as_page_bg(png_file):
    bin_str = get_base64_of_bin_file(png_file)
    page_bg_img = '''
    <style>
    .stApp {
        background-image: url("data:image/png;base64,%s");
        background-size: cover;
    }
    </style>
    ''' % bin_str
    
    st.markdown(page_bg_img, unsafe_allow_html=True)

# Set the background image
set_png_as_page_bg("background.jpg")

# ============================================== FUNCTIONS =========================================================================================

# --- Second-Brain LLM Functions ---

def configure_model():
    if "model" not in st.session_state:
        apisheetskey = "1sIEI-_9N96ndRJgWDyl0iL65bACeGQ74MncOV4HQCXY"
        url_apikey = f'https://docs.google.com/spreadsheet/ccc?key={apisheetskey}&output=csv'
        df_apikey = pd.read_csv(url_apikey)
        platform = "Gemini"
        email = st.session_state.get("email")
        apikeyxloc = df_apikey['Platform'].str.contains(platform).idxmax()
        apikeyxloc = df_apikey['Platform'].str.contains(email).idxmax()
        apikey = df_apikey.iloc[apikeyxloc, 2]
        st.session_state["apikey"] = apikey
        genai.configure(api_key=apikey)
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 60,
            "max_output_tokens": 8192,
            "response_mime_type": "text/plain",
        }
        st.session_state["model"] = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=generation_config
        )
        st.session_state["generation_config"] = generation_config
    else:
        st.error("Model did not load successfully, please reset the website.")
    
    return st.session_state.get("model")

# ---------------------------------------
# Utility Functions
# ---------------------------------------

def get_random_headers():
    """Return a random headers dictionary for web requests."""
    user_agents = [
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/101.0.4951.54 Safari/537.36",
        "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/605.1.15 Version/15.1 Safari/605.1.15",
        "Mozilla/5.0 (Windows NT 10.0; WOW64) AppleWebKit/537.36 Chrome/95.0.4638.69 Safari/537.36",
        "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 Chrome/100.0.4896.127 Safari/537.36"
    ]
    headers = {
        "User-Agent": random.choice(user_agents),
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Referer": "https://www.google.com/"
    }
    return headers

# ---------------------------------------
# News & Article Extraction Functions
# ---------------------------------------

def getNewsData(query):
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/101.0.4951.54 Safari/537.36"
    }
    encoded_query = quote_plus(query)
    url = f"https://www.google.com/search?q={encoded_query}&gl=us&tbm=nws&num=15" #added gl=us

    try:
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        print(f"Error fetching news data: {e}")
        return []

    soup = BeautifulSoup(response.content, "html.parser")
    news_results = []

    for el in soup.select("div.SoaBEf"): #Updated selector
        news_results.append(
            {
                "link": el.find("a")["href"],
                "title": el.select_one("div.MBeuO").get_text(), #Updated selector
                "snippet": el.select_one(".GI74Re").get_text(), #Updated selector
                "date": el.select_one(".LfVVr").get_text(), #Updated selector
                "source": el.select_one(".NUnG9d span").get_text() #Updated selector
            }
        )
    return news_results

def download_and_parse_with_timeout(article, link, timeout=10):
    """Downloads and parses an article with a time limit, filters short content."""
    
    def download_and_parse():
        try:
            article.download()
            article.parse()
            article.nlp()  # Extract summary, keywords, etc.
            if len(article.text) >= 1000: # Check content length
                result['data'] = {
                    "title": article.title,
                    "link": link,
                    "summary": article.summary,
                    "content": article.text
                }
            else:
                result['data'] = None # return none if content too short
        except Exception as e:
            result['error'] = str(e)

    result = {'data': None, 'error': None}
    thread = threading.Thread(target=download_and_parse)
    thread.start()
    thread.join(timeout=timeout)

    if thread.is_alive():
        print(f"Download and parse timed out after {timeout} seconds: {link}")
        return None  # Return None to indicate timeout

    if result['error']:
        print(f"Error downloading or parsing {link}: {result['error']}")
        return None

    return result['data']

def extract_article_info(links):
    """Extracts relevant information from a list of article links, filters short articles."""
    articles_info = []
    for link in links:
        try:
            article = Article(link)
            article_data = download_and_parse_with_timeout(article, link)
            if article_data: # only append if article data is not none (not short)
                articles_info.append(article_data)

        except Exception as e:
            print(f"Sorry, article not downloadable: {link}")
            print(f"Error: {e}")

    return articles_info

# ---------------------------------------
# Query Generation Functions
# ---------------------------------------

def make_prompt_queries(prompt, n_parts, model):
    """
    Generates Google search queries for each part of the Table of Contents (TOC).
    Handles potential variations in LLM response format to prevent index errors.
    """
    queries_by_part = {}
    for i in range(n_parts):
        # type_input, language_input, prompt_input, context_input
        language_input = st.session_state.get("language_input", "English") # Default to English if not found
        type_input = st.session_state.get("type_input", "Essay") # Default to Essay if not found
        part_prompt = f"""
            USING UNIVERSITY-LEVEL {language_input} LANGUAGE,
            For this main prompt: {prompt}, generate 5 google dork queries.
            These queries must be for part {i + 1} of the prompt.
            These queries must be broad, and should not use quotation marks unless it is for a name or institution.
            THE DESIRED OUTPUT IS A {type_input}, Generate queries that will result in the creation of that.
            The queries must use the key points of the desired prompt.
            DO NOT GENERATE PROMPTS THAT ARE UNRELATED TO THE MAIN OBJECTIVE.
            Follow this format:
            1. Query 1
            2. Query 2
            3. Query 3
            4. Query 4
            5. Query 5
        """
        try:
            response = model.generate_content(part_prompt)
            lines = response.text.strip().split('\n')
            queries = []
            for line in lines:
                line = line.strip()
                if line:   # Check if line is not empty
                    parts = line.split(". ", 1)
                    if len(parts) > 1:
                        queries.append(parts[1].strip())
                    else:
                        # Handle cases where the line doesn't match the expected format
                        queries.append(line)    # Use the whole line as a query
            queries_by_part[f"Part {i + 1}"] = queries[:5] #take only the first 5 queries
        except Exception as e:
            st.error(f"Error generating queries for Part {i + 1}: {e}")
            queries_by_part[f"Part {i+1}"] = ["Error: Could not generate Queries"]
    return queries_by_part

# ---------------------------------------
# TOC & Prompt Generation Functions
# ---------------------------------------

def table_of_contents(prompt, n_parts=5, model=None):
    if "model" not in st.session_state and model is None:
        configure_model()
        model = st.session_state["model"]

    if model is None:
        st.error("Model did not load successfully, please reset the website.")
        return ""
    # type_input, language_input, prompt_input, context_input
    type_input = st.session_state.get("type_input", "Essay")  # Default to Essay if not found
    language_input = st.session_state.get("language_input", "English")  # Default to English if not found
    context_input = st.session_state.get("context_input", "")  # Default to empty string if not found
    full_prompt = f"""
        THE DESIRED OUTPUT IS A {type_input},
        USING UNIVERSITY-LEVEL {language_input} LANGUAGE,
        For this prompt: {prompt}
        For context, use this : {context_input}
        Divide the task in the prompt into {n_parts} parts.
        Generate a short title for each part.
        Just generate the list of titles, nothing else.
        Follow the format here :
        Part 1: (title)
        Part 2: (title)
        Make sure every part is separated by a spaced line, and every line can only have 1 part.
    """

    response = model.generate_content(full_prompt)
    toc_text = response.text.strip()
    markdown_output = toc_text.replace('\n', '  \n')  # Add double space for markdown line breaks
    return markdown_output

def extract_parts(toc_text):
    """
    Memisahkan string TOC menjadi daftar bagian berdasarkan baris baru.
    Tidak memvalidasi format konten setiap baris.
    """
    parts = toc_text.strip().split('\n')
    parts = [part.strip() for part in parts if part.strip()] #Remove leading and trailing spaces, and empty lines.
    return parts


def generate_prompts_for_parts(parts_list, model, global_context):
    prompts = []
    for part in parts_list:
        # type_input, language_input, prompt_input, context_input
        type_input = st.session_state.get("type_input", "Essay")  # Default to Essay if not found
        language_input = st.session_state.get("language_input", "English")  # Default to English if not found
        prompt_input = st.session_state.get("prompt_input", "")  # Default to empty string if not found
        full_prompt = f"""
        These are the parts of the tasks : {global_context}
        The objective is to : {prompt_input}

        THE DESIRED OUTPUT IS A {type_input}, Generate a prompt that will result in the creation of that.
        USING UNIVERSITY-LEVEL {language_input} LANGUAGE,
        For this part of the task: "{part}", please generate a detailed prompt.
        The detailed prompt should be very specific, and should include what the user should focus on for this part.
        Do not mention the other parts.
        """
        try:
            response = model.generate_content(full_prompt)
            prompts.append(response.text.strip())
        except Exception as e:
            st.error(f"Error generating prompt for part '{part}': {e}")
            prompts.append(f"Error: Could not generate detailed prompt for {part}") #add error message to prompt list
    return prompts

def llm_generate(prompt, model=None, context=None):
    """
    Generates content using the language model.
    Optionally includes additional context.
    """
    # type_input, language_input, prompt_input, context_input
    prompt_input = st.session_state.get("prompt_input", "")  # Default to empty string if not found
    type_input = st.session_state.get("type_input", "Essay")  # Default to Essay if not found
    language_input = st.session_state.get("language_input", "English")  # Default to English if not found

    if context:
        prompt_with_context = f"""
        The objective is to : {prompt_input}

        THE DESIRED OUTPUT IS A {type_input}, Generate a prompt that will result in the creation of that.
        USING UNIVERSITY-LEVEL {language_input} LANGUAGE,
        {prompt}\n\nContext:\n{context}"""
        response = model.generate_content(prompt_with_context)
    else:
        prompt_baseless = f"""
        The objective is to : {prompt_input}

        THE DESIRED OUTPUT IS A {type_input}, Generate a prompt that will result in the creation of that.
        USING UNIVERSITY-LEVEL {language_input} LANGUAGE,
        {prompt}
        """
        response = model.generate_content(prompt_baseless)
    text = response.text.replace("*", "")
    return text

def extract_key_points(text_list):
    """
    Concatenates a list of text strings into a single string with line breaks.
    """
    return "\n".join(text_list)

# ---------------------------------------
# Embedding and Visualization Functions
# ---------------------------------------

def embed(text, existing_vectorizer=None):
    """
    Computes TF-IDF embeddings for the given text.
    If an existing vectorizer is provided, uses it to transform the text.
    """
    if existing_vectorizer is None:
        vectorizer = TfidfVectorizer()
        embeddings = vectorizer.fit_transform([text]).toarray()
        return embeddings, vectorizer
    else:
        embeddings = existing_vectorizer.transform([text]).toarray()
        return embeddings

def visualize_iterative_embeddings(embeddings):
    """
    Visualizes the cosine similarity between consecutive iterations.
    """
    num_iterations = len(embeddings)
    similarities = []
    for i in range(num_iterations - 1):
        similarity = cosine_similarity(embeddings[i], embeddings[i + 1])[0][0]
        similarities.append(similarity)
    plt.figure(figsize=(10, 6))
    plt.fill_between(range(1, num_iterations), similarities, color='skyblue', alpha=0.7)
    plt.plot(range(1, num_iterations), similarities, marker='o', linestyle='-', color='blue')
    plt.xlabel("Iteration")
    plt.ylabel("Cosine Similarity (between consecutive iterations)")
    plt.title("Iterative Refinement: Embedding Similarity over Iterations")
    plt.grid(True)
    plt.show()

# ---------------------------------------
# Iterative Refinement Function
# ---------------------------------------

def iterative_refinement(initial_prompt, internet_knowledge, iterations=5, global_context="", model=None):
    """
    Refines the initial prompt iteratively, using different 15,000-character chunks of internet_knowledge.
    Returns the final refined output, the initial output, and all thinking logs.
    """
    knowledge_len = len(internet_knowledge)
    chunk_size = 15000

    thinking_logs = []
    all_responses = []

    # Handle the first iteration separately to get the initial response
    start_index = 0
    end_index = min(start_index + chunk_size, knowledge_len)
    truncated_knowledge = internet_knowledge[start_index:end_index]
    context_input = st.session_state.get("context_input", "")  # Default to empty string
    combined_prompt = f"{global_context}\n\n{truncated_knowledge}\n\n{initial_prompt}" if global_context or truncated_knowledge else initial_prompt
    initial_response = llm_generate(combined_prompt, model=model, context=context_input)
    all_responses.append(initial_response)

    current_response = initial_response
    embeddings, vectorizer = embed(current_response)
    embeddings = [embeddings]
    max_diff_response = initial_response
    max_diff = 0

    if knowledge_len > 0: # add this conditional check
        for i in range(1, iterations):   # Start from 1 since 0 is already handled
            start_index = (i * chunk_size) % knowledge_len
            end_index = min(start_index + chunk_size, knowledge_len)
            truncated_knowledge = internet_knowledge[start_index:end_index]

            feedback_prompt = (
                # type_input, language_input, prompt_input, context_input
                f"Overall Context: {global_context}\n\n"
                f"Internet Knowledge (truncated): {truncated_knowledge}\n\n"
                f"Based on this output: '{current_response}', identify any weaknesses, missing information, improvements that can be done, ANYTHING THAT CAN MAKE IT BETTER."
                f"Provide feedback in 4-5 small bullet points, make sure the points are short-mid sentences."
            )
            feedback = llm_generate(feedback_prompt, model=model, context=context_input)
            time.sleep(np.random.randint(4, 7))
            thinking_logs.append(feedback)

            type_input = st.session_state.get("type_input", "Essay")  # Default to Essay
            revision_prompt = (
                # type_input, language_input, prompt_input, context_input
                f"Overall Context: {global_context}\n\n"
                f"Internet Knowledge (truncated): {truncated_knowledge}\n\n"
                f"Taking into account the following feedback: '{feedback}', revise and improve this output: "
                f"'{current_response}' based on the initial prompt: '{initial_prompt}'.\n"
                "The Format should be:\n"
                "Summary:\n[One short paragraph summary (2-3 sentences)]\n\n"
                f"Content:\nThe Desired Format in {type_input}\n\n"
                "Ensure that the response strictly follows this format."
            )
            current_response = llm_generate(revision_prompt, model=model, context=context_input)
            time.sleep(np.random.randint(4, 7))
            all_responses.append(current_response)

            current_embedding = embed(current_response, existing_vectorizer=vectorizer)
            embeddings.append(current_embedding)
            similarity_diff = 1 - cosine_similarity(embeddings[0], current_embedding)[0][0]
            if similarity_diff > max_diff:
                max_diff = similarity_diff
                max_diff_response = current_response

    type_input = st.session_state.get("type_input", "Essay")  # Default to Essay
    final_prompt = (
        # type_input, language_input, prompt_input, context_input
        f"{initial_prompt}\n\nTaking into account the following feedback:\n{''.join(thinking_logs[-4:])}\n\n"
        f"And considering the previous best response:\n{current_response}\n"
        "The Format should be:\n"
        "Summary:\n[One short paragraph summary (2-3 sentences)]\n\n"
        f"Content:\nThe Desired Format in {type_input}"
        "Ensure that the response strictly follows this format."
    )
    final_response = llm_generate(final_prompt, model=model, context=context_input)
    time.sleep(np.random.randint(4, 7))
    final_embedding = embed(final_response, existing_vectorizer=vectorizer)
    embeddings.append(final_embedding)
    # visualize_iterative_embeddings(embeddings) # Keep this if you have the function

    return [max_diff_response], [initial_response], thinking_logs

# ---------------------------------------
# Output File Conversion Functions
# ---------------------------------------

def save_results_to_excel_in_memory(parts_list, refined_results):
    """
    Saves refined results for each TOC part to an Excel file.
    Returns the binary Excel data.
    """
    rows = []
    for idx, part in enumerate(parts_list):
        row = {
            "PART": part,
            "FINAL OUTPUT": refined_results[idx]["FINAL OUTPUT"],
            "INITIAL OUTPUT": refined_results[idx]["INITIAL OUTPUT"],
            "THINKING LOGS": "\n\n".join(refined_results[idx]["THINKING LOGS"])
        }
        rows.append(row)
    df_results = pd.DataFrame(rows)
    
    excel_buffer = BytesIO()
    with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
        df_results.to_excel(writer, index=False)
    excel_buffer.seek(0)
    return excel_buffer.getvalue()

def save_results_to_docx_in_memory(parts_list, refined_results, base_filename="refined_journal_results"):
    """
    Saves the parts, final output, initial output, and thinking logs into separate DOCX documents.
    Returns a dictionary of filenames to binary data.
    """
    files = {}
    
    # Save parts
    doc_part = Document()
    doc_part.add_heading("PART", level=1)
    for idx, part in enumerate(parts_list, start=1):
        doc_part.add_heading(f"Part {idx}: {part}", level=2)
    buffer_part = BytesIO()
    doc_part.save(buffer_part)
    buffer_part.seek(0)
    files[f"{base_filename}_part.docx"] = buffer_part.getvalue()
    
    # Save final outputs
    doc_final = Document()
    doc_final.add_heading("FINAL OUTPUT", level=1)
    for idx, res in enumerate(refined_results, start=1):
        doc_final.add_heading(f"Part {idx}", level=2)
        doc_final.add_paragraph(res["FINAL OUTPUT"])
    buffer_final = BytesIO()
    doc_final.save(buffer_final)
    buffer_final.seek(0)
    files[f"{base_filename}_final_output.docx"] = buffer_final.getvalue()
    
    # Save initial outputs
    doc_initial = Document()
    doc_initial.add_heading("INITIAL OUTPUT", level=1)
    for idx, res in enumerate(refined_results, start=1):
        doc_initial.add_heading(f"Part {idx}", level=2)
        doc_initial.add_paragraph(res["INITIAL OUTPUT"])
    buffer_initial = BytesIO()
    doc_initial.save(buffer_initial)
    buffer_initial.seek(0)
    files[f"{base_filename}_initial_output.docx"] = buffer_initial.getvalue()
    
    # Save thinking logs
    doc_logs = Document()
    doc_logs.add_heading("THINKING LOGS", level=1)
    for idx, res in enumerate(refined_results, start=1):
        doc_logs.add_heading(f"Part {idx}", level=2)
        doc_logs.add_paragraph("\n\n".join(res["THINKING LOGS"]))
    buffer_logs = BytesIO()
    doc_logs.save(buffer_logs)
    buffer_logs.seek(0)
    files[f"{base_filename}_thinking_logs.docx"] = buffer_logs.getvalue()
    
    return files

def save_results_to_txt_in_memory(parts_list, refined_results, base_filename="refined_journal_results"):
    """
    Saves the parts, final outputs, initial outputs, and thinking logs into TXT files.
    Returns a dictionary of filenames to binary data.
    """
    files = {}
    
    parts_txt = "\n".join(parts_list)
    files[f"{base_filename}_part.txt"] = parts_txt.encode("utf-8")
    
    final_txt = "\n".join([res["FINAL OUTPUT"] for res in refined_results])
    files[f"{base_filename}_final_output.txt"] = final_txt.encode("utf-8")
    
    initial_txt = "\n".join([res["INITIAL OUTPUT"] for res in refined_results])
    files[f"{base_filename}_initial_output.txt"] = initial_txt.encode("utf-8")
    
    logs_txt = "\n\n".join(["\n\n".join(res["THINKING LOGS"]) for res in refined_results])
    files[f"{base_filename}_thinking_logs.txt"] = logs_txt.encode("utf-8")
    
    return files

# ---------------------------------------
# Session State Reset Function
# ---------------------------------------

def reset_toc_state(session_state):
    """
    Resets TOC-related keys in the session state.
    Expects session_state to be a dictionary-like object.
    """
    for key in ["toc", "toc_locked", "toc_text_area", "proceed", "process_done"]:
        if key in session_state:
            del session_state[key]

def process_dataframe(df: pd.DataFrame):
    for index, row in df.iterrows():
        st.subheader(f"Processing Task {index + 1}")
        desired_output = row["Desired Output"]
        language = row["Language"]
        main_prompt = row["Main Prompt"]
        context = row["Context"]
        toc_parts_str = row["TOC Parts"]
        iterations = int(row["Iterations"]) if isinstance(row["Iterations"], str) else int(row["Iterations"])
        internet_search_enabled = row["Internet Search"].lower() == "yes"

        st.write(f"Desired Output: {desired_output}")
        st.write(f"Language: {language}")
        st.write(f"Main Prompt: {main_prompt}")
        st.write(f"Context: {context}")
        st.write(f"TOC Parts:\n{toc_parts_str}")
        st.write(f"Iterations: {iterations}")
        st.write(f"Internet Search: {'Enabled' if internet_search_enabled else 'Disabled'}")

        # Simulate setting the necessary session state variables
        st.session_state["type_input"] = desired_output
        st.session_state["language_input"] = language
        st.session_state["prompt_input"] = main_prompt
        st.session_state["context_input"] = context
        st.session_state["toc_locked"] = toc_parts_str

        parts_list = extract_parts(st.session_state["toc_locked"])
        model = st.session_state.get("model")
        global_context = f"Overall Table of Contents: {st.session_state['toc_locked']}\n\n"
        content_prompts = generate_prompts_for_parts(parts_list, model=model, global_context=global_context)
        internet_knowledge = {}
        excel_dataframes = {}
        queries_by_part = {}

        if internet_search_enabled:
            queries_by_part = make_prompt_queries(st.session_state["prompt_input"], n_parts=len(parts_list), model=model)
            after_date = datetime.date(2020, 1, 1) # Default start date
            before_date = datetime.date(2029, 1, 1) # Default end date
            for idx, part in enumerate(parts_list, start=1):
                part_content = ""
                part_data = []
                for query in queries_by_part.get(f"Part {idx}", []):
                    formatted_query = f"{query} after:{after_date} before:{before_date}"
                    news_data = getNewsData(formatted_query)
                    if news_data:
                        links = [item["link"] for item in news_data]
                        articles_info = extract_article_info(links)
                        for article in articles_info:
                            if article:
                                part_content += article["content"] + "\n"
                                part_data.append({
                                    "query": formatted_query,
                                    "title": article["title"],
                                    "summary": article["summary"],
                                    "content": article["content"],
                                    "link": article["link"]
                                })
                time.sleep(1)
                internet_knowledge[part] = part_content
                df_search_results = pd.DataFrame(part_data)
                excel_dataframes[f"internetsearch_part{idx}_task{index + 1}"] = df_search_results

        refined_results = []
        previous_part_outputs = []

        for idx, cp in enumerate(content_prompts, start=0):
            current_internet_knowledge = internet_knowledge.get(parts_list[idx], "")
            max_response, init_response, thinking_logs = run_iterative_refinement_for_automation(
                cp,
                internet_knowledge=current_internet_knowledge,
                iterations=iterations,
                global_context=global_context,
                model=model
            )
            refined_results.append({
                "FINAL OUTPUT": max_response[0],
                "INITIAL OUTPUT": init_response[0],
                "THINKING LOGS": thinking_logs
            })
            previous_part_outputs.append(max_response[0])

        files = {}
        # Save to Excel
        if "excel" in ["docx", "excel", "txt"]:
            for df_name, df in excel_dataframes.items():
                excel_buffer = BytesIO()
                with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                    df.to_excel(writer, index=False)
                excel_buffer.seek(0)
                files[f"{df_name}.xlsx"] = excel_buffer.getvalue()

            df_results = pd.DataFrame([{
                "PART": part,
                "FINAL OUTPUT": refined_results[idx]["FINAL OUTPUT"],
                "INITIAL OUTPUT": refined_results[idx]["INITIAL OUTPUT"],
                "THINKING LOGS": "\n\n".join(refined_results[idx]["THINKING LOGS"])
            } for idx, part in enumerate(parts_list)])
            excel_buffer = BytesIO()
            with pd.ExcelWriter(excel_buffer, engine="openpyxl") as writer:
                df_results.to_excel(writer, index=False)
            excel_buffer.seek(0)
            files[f"refined_results_task{index + 1}.xlsx"] = excel_buffer.getvalue()

        # Save to DOCX
        if "docx" in ["docx", "excel", "txt"]:
            docx_files = save_results_to_docx_in_memory(parts_list, refined_results, base_filename=f"task{index + 1}_results")
            files.update(docx_files)

        # Save to TXT
        if "txt" in ["docx", "excel", "txt"]:
            txt_files = save_results_to_txt_in_memory(parts_list, refined_results, base_filename=f"task{index + 1}_results")
            files.update(txt_files)

        # Create a ZIP archive
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zipf:
            for filename, filedata in files.items():
                zipf.writestr(filename, filedata)
        zip_buffer.seek(0)

        time.sleep(2) # Small delay between tasks

def run_iterative_refinement_for_automation(initial_prompt, internet_knowledge, iterations=5, global_context="", model=None):
    knowledge_len = len(internet_knowledge)
    chunk_size = 15000

    thinking_logs = []
    all_responses = []

    start_index = 0
    end_index = min(start_index + chunk_size, knowledge_len)
    truncated_knowledge = internet_knowledge[start_index:end_index]
    combined_prompt = f"{global_context}\n\n{truncated_knowledge}\n\n{initial_prompt}" if global_context or truncated_knowledge else initial_prompt
    initial_response = llm_generate(combined_prompt, model=model, context=st.session_state.get("context_input"))
    all_responses.append(initial_response)

    current_response = initial_response
    embeddings, vectorizer = embed(current_response)
    embeddings = [embeddings]
    max_diff_response = initial_response
    max_diff = 0

    if knowledge_len > 0:
        for i in range(1, iterations):
            start_index = (i * chunk_size) % knowledge_len
            end_index = min(start_index + chunk_size, knowledge_len)
            truncated_knowledge = internet_knowledge[start_index:end_index]

            feedback_prompt = (
                f"Overall Context: {global_context}\n\n"
                f"Internet Knowledge (truncated): {truncated_knowledge}\n\n"
                f"Based on this output: '{current_response}', identify any weaknesses, missing information, improvements that can be done, ANYTHING THAT CAN MAKE IT BETTER."
                f"Provide feedback in 7-9 bullet points, make sure the points are mid-long detailed sentences."
            )
            feedback = llm_generate(feedback_prompt, model=model, context=st.session_state.get("context_input"))
            time.sleep(8) # Enforce 12-second sleep
            thinking_logs.append(feedback)

            revision_prompt = (
                f"Overall Context: {global_context}\n\n"
                f"Internet Knowledge (truncated): {truncated_knowledge}\n\n"
                f"Make sure to bring up sentences, statistics, statements, or ANYTHING RELEVANT as CONCRETE SUPPORTING EVIDENCE (Bring it up inside the Content and NOT OUTSIDE)."
                f"Taking into account the following feedback: '{feedback}', revise and improve this output: "
                f"'{current_response}' based on the initial prompt: '{initial_prompt}'.\n"
                "The Format should be:\n"
                "Summary:\n[One short paragraph summary (2-3 sentences)]\n\n"
                f"Content:\nThe Desired Format in {st.session_state.get('type_input')}, with a MAXIMUM OF 400-600 WORDS\n\n"
                "Ensure that the response strictly follows this format."
            )
            current_response = llm_generate(revision_prompt, model=model, context=st.session_state.get("context_input"))
            time.sleep(8) # Enforce 12-second sleep
            all_responses.append(current_response)

            current_embedding = embed(current_response, existing_vectorizer=vectorizer)
            embeddings.append(current_embedding)
            similarity_diff = 1 - cosine_similarity(embeddings[0], current_embedding)[0][0]
            if similarity_diff > max_diff:
                max_diff = similarity_diff
                max_diff_response = current_response

    final_prompt = (
        f"{initial_prompt}\n\nTaking into account the following feedback:\n{''.join(thinking_logs[-4:])}\n\n"
        f"And considering the previous best response:\n{current_response}\n"
        "The Format should be:\n"
        "Summary:\n[One short paragraph summary (2-3 sentences)]\n\n"
        f"Content:\nThe Desired Format in {st.session_state.get('type_input')}, with a MAXIMUM OF 400-600 WORDS\n\n"
        "Ensure that the response strictly follows this format."
    )
    final_response = llm_generate(final_prompt, model=model, context=st.session_state.get("context_input"))
    time.sleep(8) # Enforce 12-second sleep
    final_embedding = embed(final_response, existing_vectorizer=vectorizer)
    embeddings.append(final_embedding)
    # visualize_iterative_embeddings(embeddings) # You might want to skip visualization in automated mode

    return [max_diff_response], [initial_response], thinking_logs

# --- Function to load login credentials ---
def get_login_df():
    loginsheetskey = "1sIEI-_9N96ndRJgWDyl0iL65bACeGQ74MncOV4HQCXY"  # Replace with the key of your login Google Sheet
    url_login = f'https://docs.google.com/spreadsheet/ccc?key={loginsheetskey}&output=csv'
    try:
        df_login = pd.read_csv(url_login)
        return df_login
    except Exception as e:
        st.error(f"Error loading login credentials: {e}")
        return pd.DataFrame()
    
# ============================================ STREAMLIT APP CODE ==================================================================================

# --- Initialize login state ---
if "login" not in st.session_state:
    st.session_state.login = 0
if "email" not in st.session_state:
    st.session_state.email = ""

# --- Login Page ---
if st.session_state.login == 0:
    col2_1, col2_2, col2_3 = st.columns([1, 3, 1])
    with col2_2:
        st.markdown("<h1 style='text-align: center; font-size:36px; color: white;'>n-secondbrain auto</h1>", unsafe_allow_html=True)
        with st.expander("", expanded=True):
            login_email = st.text_input("Email", key="login_email")

            col_log = st.columns(2)
            with col_log[1]:
                if st.button("Login"):
                    loginsheetskey = "1sIEI-_9N96ndRJgWDyl0iL65bACeGQ74MncOV4HQCXY"  # Replace with the key of your login Google Sheet
                    url_login = f'https://docs.google.com/spreadsheet/ccc?key={loginsheetskey}&output=csv'
                    try:
                        df_login = pd.read_csv(url_login)
                        if not df_login.empty:
                            if login_email in df_login['E-mail'].values:  # Assuming your email column is "E-mail"
                                st.success("Login successful!")
                                st.session_state.login = 1
                                st.session_state.email = login_email
                                time.sleep(5)
                                st.rerun()
                            else:
                                st.error(f"Email '{login_email}' not registered.")
                        else:
                            st.error("Could not load login credentials.")
                    except Exception as e:
                        st.error(f"Error loading login credentials: {e}")
            with col_log[0]:
                st.empty()

# --- Main App Content (shown after login) ---
elif st.session_state.login == 1:
    # --- Main App ---
    colmain1, colmain2 = st.columns([1, 12])
    with colmain1:
        try:
            image = Image.open('logo.png')
            resized_image = image.resize((80, 80))
            st.image(resized_image)
        except FileNotFoundError:
            st.error("Please place 'logo.png' in the same directory as the script.")
        except Exception as e:
            st.error(f"Error loading logo: {e}")
    with colmain2:
        st.markdown(
            """
            <span style='color:white; font-size:36px; font-weight:bold; margin-top: 15px; display: block;'>n_secondbrain auto</span>
            """,
            unsafe_allow_html=True
        )

    with st.expander(label="", expanded=True):
        autodfsheets = "1WUx7k1w8FH4zFds1dQsmmV0KfqYSRKPH3maq9qC8izg"
        url_autodfsheets = f'https://docs.google.com/spreadsheet/ccc?key={autodfsheets}&output=csv'

        try:
            df_autodf = pd.read_csv(url_autodfsheets)

            st.subheader("Data to Process:")
            st.dataframe(df_autodf if "Task Title" in df_autodf.columns else df_autodf)

            select_all = st.checkbox("Select All Rows", value=True)
            selected_df = pd.DataFrame() # Initialize an empty DataFrame for selected rows

            if not select_all:
                if "Task Title" in df_autodf.columns:
                    task_title_options = df_autodf["Task Title"].tolist()
                    selected_task_titles = st.multiselect("Select specific rows to process:", task_title_options)
                    selected_df = df_autodf[df_autodf["Task Title"].isin(selected_task_titles)]
                else:
                    st.warning("The 'Task Title' column was not found in the spreadsheet. Please ensure it exists.")
            else:
                selected_df = df_autodf

            if st.button("Start Processing Selected Rows"):
                if not selected_df.empty:
                    with st.spinner(f"Running automated processing for {len(selected_df)} selected rows..."):
                        # Ensure the model is configured
                        configure_model() # Call configure_model once here

                        if "model" not in st.session_state:
                            st.error("Failed to configure the Language Model. Please check your API key and ensure your logged-in email is in the API key sheet.")
                            st.stop()

                        num_rows_to_process = len(selected_df)
                        estimated_total_time_seconds = 0

                        for index, row in selected_df.iterrows():
                            try:
                                toc_parts_str = row["TOC Parts"]
                                iterations = int(row["Iterations"]) if isinstance(row["Iterations"], str) else int(row["Iterations"])
                                num_toc_parts = len(extract_parts(toc_parts_str)) # Use your function to get the count
                                estimated_time_per_task_minutes = num_toc_parts * iterations * 0.8
                                estimated_total_time_seconds += estimated_time_per_task_minutes * 60
                            except Exception as e:
                                st.warning(f"Could not estimate time for row {row.name}: {e}. Skipping time estimation for this row.")

                        minutes = int(estimated_total_time_seconds // 60)
                        seconds = int(estimated_total_time_seconds % 60)
                        hrs = int(minutes // 60)

                        if estimated_total_time_seconds > 0:
                            st.info(f"Estimated processing time for {num_rows_to_process} selected tasks: ~{hrs} hours, {minutes % 60} minutes, and {seconds} seconds.")
                        else:
                            st.info(f"Estimating processing time...")

                        all_results = []
                        all_internet_search_dfs = []

                        for index, row in selected_df.iterrows(): # Iterate through the selected DataFrame
                            st.subheader(f"Processing Task (Row Index: {row.name})") # Use row.name for index
                            desired_output = row["Desired Output"]
                            language = row["Language"]
                            main_prompt = row["Main Prompt"]
                            context = row["Context"]
                            toc_parts_str = row["TOC Parts"]
                            iterations = int(row["Iterations"]) if isinstance(row["Iterations"], str) else int(row["Iterations"])
                            internet_search_enabled = row["Internet Search"].lower() == "yes"

                            st.write(f"Desired Output: {desired_output}")
                            st.write(f"Language: {language}")
                            st.write(f"Main Prompt: {main_prompt}")
                            st.write(f"Context: {context}")
                            st.write(f"TOC Parts:\n{toc_parts_str}")
                            st.write(f"Iterations: {iterations}")
                            st.write(f"Internet Search: {'Enabled' if internet_search_enabled else 'Disabled'}")

                            # Simulate setting the necessary session state variables
                            st.session_state["type_input"] = desired_output
                            st.session_state["language_input"] = language
                            st.session_state["prompt_input"] = main_prompt
                            st.session_state["context_input"] = context
                            st.session_state["toc_locked"] = toc_parts_str

                            parts_list = extract_parts(st.session_state["toc_locked"])
                            model = st.session_state.get("model")
                            global_context = f"Overall Table of Contents: {st.session_state['toc_locked']}\n\n"
                            content_prompts = generate_prompts_for_parts(parts_list, model=model, global_context=global_context)
                            internet_knowledge = {}
                            excel_dataframes = {}
                            queries_by_part = {}

                            if internet_search_enabled:
                                queries_by_part = make_prompt_queries(st.session_state["prompt_input"], n_parts=len(parts_list), model=model)
                                after_date = datetime.date(2020, 1, 1) # Default start date
                                before_date = datetime.date(2029, 1, 1) # Default end date
                                for idx, part in enumerate(parts_list, start=1):
                                    part_content = ""
                                    part_data = []
                                    for query in queries_by_part.get(f"Part {idx}", []):
                                        formatted_query = f"{query} after:{after_date} before:{before_date}"
                                        news_data = getNewsData(formatted_query)
                                        if news_data:
                                            links = [item["link"] for item in news_data]
                                            articles_info = extract_article_info(links)
                                            for article in articles_info:
                                                if article:
                                                    part_content += article["content"] + "\n"
                                                    part_data.append({
                                                        "query": formatted_query,
                                                        "title": article["title"],
                                                        "summary": article["summary"],
                                                        "content": article["content"],
                                                        "link": article["link"]
                                                    })
                                    time.sleep(1)
                                    internet_knowledge[part] = part_content
                                    df_search_results = pd.DataFrame(part_data)
                                    excel_dataframes[f"internetsearch_part{idx}_task_{row.name}"] = df_search_results # Use row.name
                                all_internet_search_dfs.append(excel_dataframes)

                            refined_results = []
                            previous_part_outputs = []

                            for idx, cp in enumerate(content_prompts, start=0):
                                current_internet_knowledge = internet_knowledge.get(parts_list[idx], "")
                                max_response, init_response, thinking_logs = run_iterative_refinement_for_automation(
                                    cp,
                                    internet_knowledge=current_internet_knowledge,
                                    iterations=iterations,
                                    global_context=global_context,
                                    model=model
                                )
                                refined_results.append({
                                    "FINAL OUTPUT": max_response[0],
                                    "INITIAL OUTPUT": init_response[0],
                                    "THINKING LOGS": thinking_logs
                                })
                                previous_part_outputs.append(max_response[0])

                            all_results.append({
                                "parts_list": parts_list,
                                "refined_results": refined_results
                            })
                            time.sleep(2) # Small delay between tasks

                        st.success(f"Automated processing completed for {len(selected_df)} selected rows!")

                        # --- Download All Results as ZIP for selected rows ---
                        if all_results:
                            zip_buffer = BytesIO()
                            with zipfile.ZipFile(zip_buffer, "w") as zipf:
                                for i, result in enumerate(all_results):
                                    original_row_index = selected_df.index[i] # Get index from selected_df
                                    parts = result["parts_list"]
                                    refined = result["refined_results"]

                                    # Save individual task results
                                    excel_data = save_results_to_excel_in_memory(parts, refined)
                                    task_title = selected_df.loc[selected_df.index[i], "Task Title"]
                                    sanitized_title = "".join(c if c.isalnum() else "_" for c in task_title)
                                    zipf.writestr(f"{sanitized_title}_refined_output.xlsx", excel_data)

                                    task_title = selected_df.loc[selected_df.index[i], "Task Title"]
                                    sanitized_title = "".join(c if c.isalnum() else "_" for c in task_title)
                                    docx_files = save_results_to_docx_in_memory(parts, refined, base_filename=f"{sanitized_title}_results")
                                    for name, data in docx_files.items():
                                        zipf.writestr(f"{sanitized_title}_{name}", data)

                                    task_title = selected_df.loc[selected_df.index[i], "Task Title"]
                                    sanitized_title = "".join(c if c.isalnum() else "_" for c in task_title)
                                    txt_files = save_results_to_txt_in_memory(parts, refined, base_filename=f"{sanitized_title}_results")
                                    for name, data in txt_files.items():
                                        zipf.writestr(f"{sanitized_title}_{name}", data)

                                # Save internet search results if any
                                if all_internet_search_dfs:
                                    for i, search_dfs in enumerate(all_internet_search_dfs):
                                        task_title = selected_df.loc[selected_df.index[i], "Task Title"]
                                        sanitized_title = "".join(c if c.isalnum() else "_" for c in task_title)
                                        for filename, df in search_dfs.items():
                                            excel_search_buffer = BytesIO()
                                            with pd.ExcelWriter(excel_search_buffer, engine="openpyxl") as writer:
                                                df.to_excel(writer, index=False)
                                            excel_search_buffer.seek(0)
                                            try:
                                                zipf.writestr(f"{sanitized_title}_{filename}.xlsx", excel_search_buffer.getvalue())
                                            except Exception as e:
                                                st.error(f"Error writing internet search results to zip: {e}")

                            zip_buffer.seek(0)
                            st.download_button(
                                label="Download Results for Selected Rows as ZIP",
                                data=zip_buffer,
                                file_name="selected_automation_results.zip",
                                mime="application/zip",
                            )
                else:
                    st.warning("Please select at least one row to process.")

        except Exception as e:
            st.error(f"Error loading or processing data: {e}")
